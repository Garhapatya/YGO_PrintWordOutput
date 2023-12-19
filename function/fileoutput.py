import docx
from docx import Document
from docx.shared import Cm

import tkinter as tk
from tkinter import messagebox

from function import card_pic

def newdocx(Cards:list,DocName:str) ->str:

    
    try:
        docu = Document(".\\resouse\\预设样板.docx")
    except docx.opc.exceptions.PackageNotFoundError:
        docu = Document()
        docu.sections[0].page_width = Cm(21)
        docu.sections[0].page_height = Cm(29.7)
        docu.sections[0].left_margin = Cm(1.27)
        docu.sections[0].right_margin = Cm(1.27)
        docu.sections[0].top_margin = Cm(0.45)
        docu.sections[0].bottom_margin = Cm(0.45)
        docu.save(".\\resouse\\预设样板.docx")

    para = docu.add_paragraph()
    run = para.add_run('\n')
    #初始化样板页面布局

    
    count=0
    for Card in Cards:
        try:
            run.add_picture(card_pic.card_address(Card[0]), width=Cm(5.9),height=Cm(8.6))
        except FileNotFoundError:
            run.add_picture(".\\resouse\\"+"missingcardpic"+".jpg", width=Cm(5.9),height=Cm(8.6))
            #卡图缺少的卡片会使用“暂无卡图”样式图片代替
        
        count+=1
        if count==3:
            count=0
            run.add_break()
        #每三张图换行



    while 1:
        try:
            docu.save(DocName+".docx")
            break
        except PermissionError:
                messagebox.showwarning("保存失败","检查并确认无其他进程占用该文件并重试")
                


#newdocx([('84206435', '冰结界的饰章'), ('43582229', '冰结界的晶壁'), ('70703416', '冰结界的霜精'), ('34293667', '冰结界'), ('44308317', '冰结界的依巫'), ('23950192', '冰结界的术者'), ('17197110', '到达冰结界的晴岚'), ('64990807', '冰结界的三方阵'), ('50032342', '冰结界的军师'), ('43256007', ' 冰结界的随身'), ('53535814', '冰结界的净玻璃'), ('81825063', '冰结界的虎将 韦恩'), ('59546528', '冰结界的舞姬'), ('18319762', '冰结界的照魔师'), ('35380371', '到达冰结界的晶域'), ('9396662', '冰结界的镜魔师'), ('82498947', '冰结界的守护阵'), ('27527047', '冰结界的御庭番'), ('50088247', '冰结界的传道师'), ('90311614', '冰结界的水影'), ('37806313', '冰结界的输送部队'), ('73061465', '冰结界的封魔团'), ('18482591', '冰结界的破术 师'), ('64926005', '冰结界的武士'), ('20700531', '冰结界的修行者'), ('38528901', '冰结界的奋起队'), ('41090784', '冰结界的大僧正'), ('53921056', '冰结界的虎将 乾陀罗'), ('32991027', '冰结界的剑士 格奥鲁吉尔斯'), ('9056100', '冰结界的虎将 格尔纳德'), ('10691144', '冰结界之镜'), ('60161788', '冰结界的看守人 布里兹德'), ('44877690', '冰结界的神精灵'), ('56704140', '冰结界的风水师'), ('92065772', '住在冰结界的魔醉虫'), ('70980824', '冰结界之还零龙 特里修拉'), ('81275309', '冰结界的虎将 莱蓬'), ('66661678', '冰结界的王家骑士'), ('88494899', '冰结界的交灵师'), ('65749035', '冰结界之龙 冈尼尔'), ('70583986', '冰结界的虎王 杜罗伦'), ('50321796', '冰结界之龙 布琉纳克'), ('52687916', '冰结界之龙 特里修拉'), ('6075533', '特里修拉的鼓动'), ('96402918', '冰灵山的龙祖 兰赛亚')])
